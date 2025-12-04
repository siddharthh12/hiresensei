import requests
import os
from docx import Document

BASE_URL = "http://localhost:8000"

def create_dummy_resume():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Email: johndoe@example.com')
    doc.add_paragraph('Phone: 123-456-7890')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, FastAPI, React, MongoDB')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Engineer at Tech Corp')
    doc.add_paragraph('Developed web applications.')
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.Sc. Computer Science')
    
    filename = "test_resume.docx"
    doc.save(filename)
    return filename

def verify_resume_flow():
    print("Starting Resume Flow Verification...")
    
    # 1. Login to get token
    print("Logging in...")
    login_data = {
        "username": "testuser@example.com",
        "password": "password123"
    }
    # Ensure user exists (re-register if needed, or just login)
    # Assuming testuser exists from previous step. If not, register.
    requests.post(f"{BASE_URL}/auth/register", json={
        "full_name": "Test User",
        "email": "testuser@example.com",
        "password": "password123"
    })
    
    res = requests.post(f"{BASE_URL}/auth/login", data=login_data)
    if res.status_code != 200:
        print(f"Login failed: {res.text}")
        return
    token = res.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    print("Login successful.")

    # 2. Create Dummy Resume
    filename = create_dummy_resume()
    print(f"Created dummy resume: {filename}")

    # 3. Upload
    print("Uploading resume...")
    with open(filename, "rb") as f:
        files = {"file": (filename, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        res = requests.post(f"{BASE_URL}/resume/upload", files=files, headers=headers)
    
    if res.status_code != 200:
        print(f"Upload failed: {res.text}")
        return
    
    file_path = res.json()["file_path"]
    print(f"Upload successful. Path: {file_path}")

    # 4. Parse
    print("Parsing resume...")
    res = requests.post(f"{BASE_URL}/resume/parse", params={"file_path": file_path}, headers=headers)
    
    if res.status_code != 200:
        print(f"Parsing failed: {res.text}")
        return
    
    parsed_data = res.json()
    print("Parsed Data:")
    print(parsed_data)
    
    # Verify extraction
    if parsed_data["email"] == "johndoe@example.com":
        print("Email extraction verified.")
    else:
        print("Email extraction failed.")

    # 5. Save
    print("Saving profile...")
    parsed_data["file_path"] = file_path
    res = requests.post(f"{BASE_URL}/resume/save", json=parsed_data, headers=headers)
    
    if res.status_code != 200:
        print(f"Save failed: {res.text}")
        return
    
    saved_profile = res.json()
    print("Profile saved successfully.")
    print(saved_profile)

    # Cleanup
    if os.path.exists(filename):
        os.remove(filename)
    # Note: Uploaded file in backend/uploads is left there.

if __name__ == "__main__":
    verify_resume_flow()
