import requests
import os
from docx import Document
import uuid

BASE_URL = "http://localhost:8000"
TEST_EMAIL = f"test_{uuid.uuid4()}@example.com"
TEST_PASSWORD = "password123"

def print_result(name, success, details=""):
    status = "✅ PASS" if success else "❌ FAIL"
    print(f"{status} - {name}")
    if not success and details:
        print(f"   Error: {details}")

def create_dummy_resume():
    filename = f"test_resume_{uuid.uuid4()}.docx"
    doc = Document()
    doc.add_heading('Jane Doe', 0)
    doc.add_paragraph(f'Email: {TEST_EMAIL}')
    doc.add_paragraph('Skills: Python, Testing')
    doc.save(filename)
    return filename

def test_api():
    print(f"Starting API Test Suite... (User: {TEST_EMAIL})")
    print("-" * 50)

    token = None
    job_id = None
    file_path = None

    # 1. Auth: Register
    try:
        res = requests.post(f"{BASE_URL}/auth/register", json={
            "full_name": "Test User",
            "email": TEST_EMAIL,
            "password": TEST_PASSWORD
        })
        print_result("POST /auth/register", res.status_code == 200, res.text)
    except Exception as e:
        print_result("POST /auth/register", False, str(e))

    # 2. Auth: Login
    try:
        res = requests.post(f"{BASE_URL}/auth/login", data={
            "username": TEST_EMAIL,
            "password": TEST_PASSWORD
        })
        if res.status_code == 200:
            token = res.json()["access_token"]
            print_result("POST /auth/login", True)
        else:
            print_result("POST /auth/login", False, res.text)
    except Exception as e:
        print_result("POST /auth/login", False, str(e))

    headers = {"Authorization": f"Bearer {token}"} if token else {}

    # 3. Auth: Me
    if token:
        try:
            res = requests.get(f"{BASE_URL}/auth/me", headers=headers)
            print_result("GET /auth/me", res.status_code == 200, res.text)
        except Exception as e:
            print_result("GET /auth/me", False, str(e))
    else:
        print_result("GET /auth/me", False, "Skipped (No Token)")

    # 4. Jobs: Create
    try:
        res = requests.post(f"{BASE_URL}/job", json={
            "title": "Test Job",
            "description": "Test Description",
            "company": "Test Co",
            "location": "Remote"
        }) # Note: Job endpoints might not be protected yet based on previous code, but let's see.
           # Checking main.py from previous turns, job endpoints are NOT protected.
        if res.status_code == 200:
            job_id = res.json()["data"][0]["id"]
            print_result("POST /job", True)
        else:
            print_result("POST /job", False, res.text)
    except Exception as e:
        print_result("POST /job", False, str(e))

    # 5. Jobs: Get All
    try:
        res = requests.get(f"{BASE_URL}/jobs")
        print_result("GET /jobs", res.status_code == 200, res.text)
    except Exception as e:
        print_result("GET /jobs", False, str(e))

    # 6. Jobs: Get One
    if job_id:
        try:
            res = requests.get(f"{BASE_URL}/job/{job_id}")
            print_result(f"GET /job/{{id}}", res.status_code == 200, res.text)
        except Exception as e:
            print_result(f"GET /job/{{id}}", False, str(e))
    else:
        print_result(f"GET /job/{{id}}", False, "Skipped (No Job ID)")

    # 7. Jobs: Update
    if job_id:
        try:
            res = requests.put(f"{BASE_URL}/job/{job_id}", json={"title": "Updated Job"})
            print_result(f"PUT /job/{{id}}", res.status_code == 200, res.text)
        except Exception as e:
            print_result(f"PUT /job/{{id}}", False, str(e))
    else:
        print_result(f"PUT /job/{{id}}", False, "Skipped (No Job ID)")

    # 8. Jobs: Delete
    if job_id:
        try:
            res = requests.delete(f"{BASE_URL}/job/{job_id}")
            print_result(f"DELETE /job/{{id}}", res.status_code == 200, res.text)
        except Exception as e:
            print_result(f"DELETE /job/{{id}}", False, str(e))
    else:
        print_result(f"DELETE /job/{{id}}", False, "Skipped (No Job ID)")

    # 9. Resume: Upload
    resume_file = create_dummy_resume()
    if token:
        try:
            with open(resume_file, "rb") as f:
                files = {"file": (resume_file, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                res = requests.post(f"{BASE_URL}/resume/upload", files=files, headers=headers)
            
            if res.status_code == 200:
                file_path = res.json()["file_path"]
                print_result("POST /resume/upload", True)
            else:
                print_result("POST /resume/upload", False, res.text)
        except Exception as e:
            print_result("POST /resume/upload", False, str(e))
    else:
        print_result("POST /resume/upload", False, "Skipped (No Token)")

    # 10. Resume: Parse
    if token and file_path:
        try:
            res = requests.post(f"{BASE_URL}/resume/parse", params={"file_path": file_path}, headers=headers)
            print_result("POST /resume/parse", res.status_code == 200, res.text)
        except Exception as e:
            print_result("POST /resume/parse", False, str(e))
    else:
        print_result("POST /resume/parse", False, "Skipped (No Token or File Path)")

    # 11. Resume: Save
    if token and file_path:
        try:
            # Need valid data structure for save
            data = {
                "name": "Jane Doe",
                "email": TEST_EMAIL,
                "phone": "1234567890",
                "skills": ["Python"],
                "experience": [],
                "education": [],
                "certifications": []
            }
            res = requests.post(f"{BASE_URL}/resume/save", params={"file_path": file_path}, json=data, headers=headers)
            print_result("POST /resume/save", res.status_code == 200, res.text)
        except Exception as e:
            print_result("POST /resume/save", False, str(e))
    else:
        print_result("POST /resume/save", False, "Skipped (No Token or File Path)")

    # Cleanup
    if os.path.exists(resume_file):
        os.remove(resume_file)

if __name__ == "__main__":
    test_api()
